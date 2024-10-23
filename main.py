import os
import streamlit as st
from PIL import Image, ImageOps
import pandas as pd
import base64
import matplotlib.pyplot as plt
from docx import Document
import zipfile
import json
import toml

# Set initial working directory
if 'current_path' not in st.session_state:
    st.session_state.current_path = "C:\\"

# Initialize navigation history
if 'nav_history' not in st.session_state:
    st.session_state.nav_history = []


# Helper function for path updates
def update_path(new_path):

    if new_path != st.session_state.current_path:
        st.session_state.nav_history.append(st.session_state.current_path)

    st.session_state.current_path = new_path


def read_docx(filepath):
    doc = Document(filepath)
    text = ""

    for para in doc.paragraphs:
        text += para.text + "\n"

    return text


if 'selected_file' not in st.session_state:
    st.session_state.selected_file = None

# Directory navigation
directory = st.text_input("Give The Path Of Directory")
if st.button("Change Directory"):
    if os.path.isdir(directory):
        update_path(directory)
        st.session_state.selected_file = None
        st.rerun()

# Breadcrumb Trail
st.markdown(f"### Current Directory: {st.session_state.current_path}")

# Split the current path into parts to create breadcrumb buttons
path_parts = st.session_state.current_path.split(os.sep)

# Construct the breadcrumb trail with buttons
path_to_root = path_parts[0] if len(path_parts) > 0 else ""

for part in path_parts[1:]:

    if part:  # Ensure the part is not empty
        path_to_root = os.path.join(path_to_root, part)

if st.button("Home"):
    update_path("C:\\")
    st.session_state.selected_file = None
    st.rerun()

# Get current working directory
path = st.session_state.current_path
files = os.listdir(path)

# Sort directories to appear at the top
files.sort(key=lambda f: (not os.path.isdir(os.path.join(path, f)), f.lower()))

# Search functionality
search_query = st.text_input("Search for files or folders:")
filtered_files = [f for f in files if search_query.lower() in f.lower()]

# If a file is currently being edited
if st.session_state.selected_file:
    file_path = st.session_state.selected_file

    if file_path.endswith((".png", ".jpg", ".jpeg", ".svg", ".gif", ".webp")):
        # Image Editor
        st.title(f"Editing Image: {os.path.basename(file_path)}")
        original, edited = st.columns(2)

        # Display the original image
        image = Image.open(file_path)
        original.image(image, caption="Original Image", use_column_width=True)

        # Rotate the image
        rotate_degrees = st.slider("Rotate Image", 0, 360, 0)
        rotated_image = image.rotate(rotate_degrees)

        # Add a checkbox to flip the image
        if st.checkbox("Flip Image Horizontally"):
            rotated_image = ImageOps.mirror(rotated_image)
        
        if st.checkbox("Flip Image Vertically"):
            rotated_image = ImageOps.flip(rotated_image)

        # Cropping
        left = st.slider("Crop Left", 0, image.width, 0)
        top = st.slider("Crop Top", 0, image.height, 0)
        right = st.slider("Crop Right", left, image.width, image.width)
        bottom = st.slider("Crop Bottom", top, image.height, image.height)
        cropped_image = rotated_image.crop((left, top, right, bottom))

        # Apply grayscale filter
        if st.checkbox("Convert to Grayscale"):
            cropped_image = ImageOps.grayscale(cropped_image)

        # Display the edited image
        edited.image(cropped_image, caption="Edited Image", use_column_width=True)

        # Button to save the edited image
        if st.button("Save Changes"):
            cropped_image.save(file_path)
            st.success("Image saved successfully")

        # Button to go back to file explorer
        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(".txt"):
        # Text Editor
        st.title(f"Editing: {os.path.basename(file_path)}")

        # Display the content of the selected file
        with open(file_path, 'r') as file:
            content = file.read()

        # Provide a text area for editing the content
        new_content = st.text_area("Edit content:", value=content, height=300)

        # Button to save the edited content
        if st.button("Save Changes"):
            with open(file_path, 'w') as file:
                file.write(new_content)
            st.success("File saved successfully")

        # Button to go back to file explorer
        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith("pdf"):
        st.title(f"Viewing: {os.path.basename(file_path)}")
        
        with open(file_path, "rb") as pdf_file:
            pdf_data = pdf_file.read()

            base64_pdf = base64.b64encode(pdf_data).decode("utf-8")
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(("mp3", "wav")):
        st.title(f"Listening: {os.path.basename(file_path)}")
        st.audio(file_path, format="audio/wav" if file_path.endswith("wav") else "audio/mpeg")

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith("mp4"):
        st.title(f"Viewing: {os.path.basename(file_path)}")
        st.video(file_path)

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(".md"):
        # Markdown Editor
        st.title(f"Editing Markdown File: {os.path.basename(file_path)}")

        # Display the content of the selected markdown file
        with open(file_path, 'r') as file:
            content = file.read()

        # Provide a text area for editing the content
        new_content = st.text_area("Edit Markdown:", value=content, height=300)
        st.markdown(new_content)  # Preview markdown

        # Button to save the edited content
        if st.button("Save Changes"):
        
            with open(file_path, 'w') as file:
                file.write(new_content)
        
            st.success("Markdown file saved successfully")

        # Button to go back to file explorer
        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(("csv", "xlsx")):
        st.title(f"Editing: {os.path.basename(file_path)}")
        
        if file_path.endswith("csv"):
            df = pd.read_csv(file_path)
            data = st.data_editor(df)

            if st.button("Save Changes"):
                data.to_csv(file_path, index=False)
                st.success("File saved successfully")

            # Button to go back to file explorer
            if st.button("Back to File Explorer"):
                st.session_state.selected_file = None
                st.rerun()

            # Visualization Options
            st.subheader("Visualize Data")
            columns = st.multiselect("Select columns to plot", df.columns)
            
            if len(columns) >= 2:
                st.line_chart(df[columns])

            if st.button("Plot with Matplotlib"):
                fig, ax = plt.subplots()
                df[columns].plot(ax=ax)
                st.pyplot(fig)

        if file_path.endswith("xlsx"):
            with pd.ExcelFile(file_path) as xls:
                sheets = xls.sheet_names
                sheet = st.radio("Choose sheet name:", sheets)
                df = pd.read_excel(file_path, sheet)
                data = st.data_editor(df)

                if st.button("Save Changes"):
            
                    with pd.ExcelWriter(file_path) as writer:
                        data.to_excel(writer, sheet_name=sheet)
            
                    st.success("File saved successfully")

                if st.button("Return To File Explorer"):
                    st.session_state.selected_file = None
                    st.rerun()

                # Visualization Options
                st.subheader("Visualize Data")
                columns = st.multiselect("Select columns to plot", df.columns)
            
                if len(columns) >= 2:
                    st.line_chart(df[columns])

                if st.button("Plot with Matplotlib"):
                    fig, ax = plt.subplots()
                    df[columns].plot(ax=ax)
                    st.pyplot(fig)

    elif file_path.endswith(".docx"):
        st.title(f"Editing: {os.path.basename(file_path)}")
        content = read_docx(file_path)
        new_content = st.text_area("Edit content:", value=content, height=300)

        if st.button("Save Changes"):
            doc = Document()
            
            for line in new_content.split("\n"):
                doc.add_paragraph(line)
            
            doc.save(file_path)
            st.success("File saved successfully")

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(".zip"):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            zip_info = zip_ref.infolist()
            st.write("Files in the ZIP archive:")

            # Display the list of files inside the ZIP archive
            for info in zip_info:
                st.text(info.filename)

            # Option to extract all files
            if st.button("Extract All Files"):
                extract_path = st.text_input("Provide the extraction path:", value=path)
                
                if extract_path and os.path.isdir(extract_path):
                    zip_ref.extractall(extract_path)
                    st.success("ZIP archive extracted successfully")
                
                else:
                    st.error("Invalid extraction path")

        # Button to go back to file explorer
        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(".ipynb"):
        # Jupyter Notebook Viewer
        st.title(f"Viewing: {os.path.basename(file_path)}")
        
        "---"

        # Display the content of the notebook in JSON format
        with open(file_path, 'r') as file:
            notebook_content = json.load(file)

        for i in range(len(notebook_content["cells"])):
            if notebook_content["cells"][i]["cell_type"] == "code":
                source = notebook_content["cells"][i]["source"]
                temp = ""
                
                for block in source:
                    temp += block
                
                st.code(temp, language="python")

            else:
                source = notebook_content["cells"][i]["source"]
                temp = ""
                
                for block in source:
                    temp += block
                
                st.markdown(temp)

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(".toml"):
        # TOML Editor
        st.title(f"Editing: {os.path.basename(file_path)}")

        # Load the content of the TOML file
        with open(file_path, 'r') as file:
            toml_content = file.read()

        # Provide a text area for editing the content
        new_toml_content = st.text_area("Edit TOML content:", value=toml_content, height=300)

        # Button to save the edited content
        if st.button("Save Changes"):
            try:
                # Validate if the new content is a valid TOML format
                toml.loads(new_toml_content)
                
                with open(file_path, 'w') as file:
                    file.write(new_toml_content)
                
                st.success("TOML file saved successfully")
            
            except toml.TomlDecodeError:
                st.error("Invalid TOML format. Please fix the errors before saving.")

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith("json"):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with open(file_path, "r") as file:
            content = json.load(file)
        
        st.json(content)
        
        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()
    
    elif file_path.endswith(("html", "css", "js", "php", "ts", "tsx")):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with open(file_path, "r") as file:
            content = file.read()
        
        st.code(content, language="html" if file_path.endswith("html") else ("css" if file_path.endswith("css") else ("js" if file_path.endswith(("js", "ts", "tsx")) else "php")))

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()
        
    elif file_path.endswith(("c", "c++", "cs")):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with open(file_path, "r") as file:
            content = file.read()
        
        st.code(content, language="c" if file_path.endswith("c") else ("c++" if file_path.endswith("c++") else "cs"))

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()
        
    elif file_path.endswith("java"):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with open(file_path, "r") as file:
            content = file.read()
        
        st.code(content, language="java")

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith("sql"):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with open(file_path, "r") as file:
            content = file.read()

        st.code(content, language="sql")

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()
        
    elif file_path.endswith("asm"):
        st.title(f"Viewing: {os.path.basename(file_path)}")

        with open(file_path, "r") as file:
            content = file.read()
        
        st.code(content, language="asm")

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    elif file_path.endswith(".py"):
        st.title(f"Viewing: {os.path.basename(file_path)}")
        with open(file_path, 'r') as file:
            content = file.read()

        st.code(content, language='python')

        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

    else:
        st.markdown("# :red[File Format Not Supported !!]")
        
        if st.button("Back to File Explorer"):
            st.session_state.selected_file = None
            st.rerun()

# If no file is selected, show the file explorer
else:
    # Create three columns
    col1, col2, col3 = st.columns(3)

    # Display files and folders as buttons in three columns
    for idx, file in enumerate(files):
        # Assign the columns dynamically for better layout
        if idx % 3 == 0:
            column = col1
        
        elif idx % 3 == 1:
            column = col2
        
        else:
            column = col3

        # Use the column to display the button
        with column:
            base_dir = os.path.dirname(__file__)
            icons_dir = os.path.join(base_dir, "icons")

            # Define paths to the icons
            txt_icon = os.path.join(icons_dir, "txt.jpg")
            python_icon = os.path.join(icons_dir, "python.jpg")
            folder_icon = os.path.join(icons_dir, "folder.jpg")
            image_icon = os.path.join(icons_dir, "image.jpg")
            docx_icon = os.path.join(icons_dir, "docx.jpg")
            exe_icon = os.path.join(icons_dir, "exe.jpg")
            pdf_icon = os.path.join(icons_dir, "pdf.jpg")
            audio_icon = os.path.join(icons_dir, "audio.jpg")
            video_icon = os.path.join(icons_dir, "video.jpg")
            excel_icon = os.path.join(icons_dir, "excel.jpg")
            file_icon = os.path.join(icons_dir, "file.jpg")
            zip_icon = os.path.join(icons_dir, "zip.jpg")
            ipynb_icon = os.path.join(icons_dir, "ipynb.jpg")
            json_icon = os.path.join(icons_dir, "json.jpg")
            html_icon = os.path.join(icons_dir, "html.jpg")
            css_icon = os.path.join(icons_dir, "css.jpg")
            js_icon = os.path.join(icons_dir, "js.jpg")
            php_icon = os.path.join(icons_dir, "php.jpg")
            c_icon = os.path.join(icons_dir, "c.jpg")
            cplusplus_icon = os.path.join(icons_dir, "c-plus-plus.jpg")
            csharp_icon = os.path.join(icons_dir, "c-sharp.jpg")
            java_icon = os.path.join(icons_dir, "java.jpg")
            assembly_icon = os.path.join(icons_dir, "assembly.jpg")
            ts_icon = os.path.join(icons_dir, "ts.jpg")
            sql_icon = os.path.join(icons_dir, "sql.jpg")
            ppt_icon = os.path.join(icons_dir, "ppt.jpg")

            # Display the appropriate icon for each file type
            if file.endswith("txt"):
                icon = txt_icon if os.path.exists(txt_icon) else None
                label = "Text File"

            elif file.endswith("py"):
                icon = python_icon if os.path.exists(python_icon) else None
                label = "Python File"

            elif file.endswith("pptx"):
                icon = ppt_icon if os.path.exists(ppt_icon) else None
                label = "Powerpoint Presentation"

            elif file.endswith(("ts", "tsx")):
                icon = ts_icon if os.path.exists(ts_icon) else None
                label = "Typescript File"

            elif file.endswith("sql"):
                icon = sql_icon if os.path.exists(sql_icon) else None
                label = "SQL File"

            elif file.endswith("zip"):
                icon = zip_icon if os.path.exists(zip_icon) else None
                label = "ZIP Archive"

            elif file.endswith("ipynb"):
                icon = ipynb_icon if os.path.exists(ipynb_icon) else None
                label = "Jupyter Notebook"

            elif os.path.isdir(os.path.join(path, file)):
                icon = folder_icon if os.path.exists(folder_icon) else None
                label = "Folder"

            elif file.endswith(("png", "jpg", "svg", "gif", "webp", "jpeg")):
                icon = image_icon if os.path.exists(image_icon) else None
                label = "Image"

            elif file.endswith("docx"):
                icon = docx_icon if os.path.exists(docx_icon) else None
                label = "Microsoft Word File"

            elif file.endswith("html"):
                icon = html_icon if os.path.exists(html_icon) else None
                label = "HTML File"
            
            elif file.endswith("css"):
                icon = css_icon if os.path.exists(css_icon) else None
                label = "CSS File"
            
            elif file.endswith("c"):
                icon = c_icon if os.path.exists(c_icon) else None
                label = "C File"
            
            elif file.endswith("c++"):
                icon = cplusplus_icon if os.path.exists(cplusplus_icon) else None
                label = "C++ File"

            elif file.endswith("cs"):
                icon = csharp_icon if os.path.exists(csharp_icon) else None
                label = "C# File"

            elif file.endswith("java"):
                icon = java_icon if os.path.exists(java_icon) else None
                label = "Java File"

            elif file.endswith("asm"):
                icon = assembly_icon if os.path.exists(assembly_icon) else None
                label = "Assembly File"

            elif file.endswith("js"):
                icon = js_icon if os.path.exists(js_icon) else None
                label = "Javascript File"
            
            elif file.endswith("php"):
                icon = php_icon if os.path.exists(php_icon) else None
                label = "PHP File"

            elif file.endswith("pdf"):
                icon = pdf_icon if os.path.exists(pdf_icon) else None
                label = "PDF"

            elif file.endswith("exe"):
                icon = exe_icon if os.path.exists(exe_icon) else None
                label = "Executable File"

            elif file.endswith(("wav", "mp3")):
                icon = audio_icon if os.path.exists(audio_icon) else None
                label = "Audio File"

            elif file.endswith("mp4"):
                icon = video_icon if os.path.exists(video_icon) else None
                label = "Video File"

            elif file.endswith(("xlsx", "csv")):
                icon = excel_icon if os.path.exists(excel_icon) else None
                label = "Excel File"
            
            elif file.endswith("json"):
                icon = json_icon if os.path.exists(json_icon) else None
                label = "Json File"

            else:
                icon = file_icon if os.path.exists(file_icon) else None
                label = "Unknown File"

            # Display icon or label
            if icon:
                st.image(icon, width=40)

            else:
                st.write(label)

            # Display the button with label
            if st.button(file, key=f"{path}_{file}"):
                new_path = os.path.join(path, file)

                if os.path.isdir(new_path):
                    st.session_state.current_path = new_path
                    st.rerun()

                elif os.path.isfile(new_path):
                    st.session_state.selected_file = new_path
                    st.rerun()

            "---"
